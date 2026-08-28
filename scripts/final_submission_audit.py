"""Final submission audit — read-only checks."""

from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


def main() -> int:
    root = PROJECT_ROOT
    results: list[tuple[str, str, str]] = []

    def add(req: str, ok: bool, evidence: str) -> None:
        results.append((req, "PASS" if ok else "FAIL", evidence))

    required_files = [
        "validation_predictions.csv",
        "december-chart-inputs.csv",
        "scorer_results/candidate_december.png",
        "README.md",
        "requirements.txt",
        "score.py",
    ]
    report_docx = root / "reports/freight_rate_ml_assessment.docx"
    report_pdf = root / "reports/freight_rate_ml_assessment.pdf"
    report_ok = report_docx.is_file() or report_pdf.is_file()
    add(
        "Final report (DOCX or PDF)",
        report_ok,
        str(report_docx if report_docx.is_file() else report_pdf),
    )

    for f in required_files:
        p = root / f
        add(f"File exists: {f}", p.is_file(), f"{p.stat().st_size} bytes" if p.is_file() else "missing")

    for doc in [
        "docs/PHASE_0_PROJECT_AUDIT.md",
        "docs/PHASE_1_EDA.md",
        "docs/PHASE_2_VALIDATION.md",
        "docs/PHASE_3_MODELING.md",
        "docs/PHASE_4_FINAL_MODEL.md",
        "docs/PHASE_5_SCORING.md",
        "docs/LOOM_SCRIPT.md",
        "docs/GITHUB_CHECKLIST.md",
    ]:
        p = root / doc
        add(f"Doc exists: {doc}", p.is_file(), "present" if p.is_file() else "missing")

    pred = pd.read_csv(root / "validation_predictions.csv")
    tmpl = pd.read_csv(root / "validation-predictions-template.csv")
    val = pd.read_csv(root / "validation.csv")
    rates = pd.to_numeric(pred["predicted_rate"], errors="coerce")

    add("Predictions: 12,000 rows", len(pred) == 12_000, str(len(pred)))
    add("Predictions: 2 columns", pred.shape[1] == 2, str(list(pred.columns)))
    add("Predictions: column names", list(pred.columns) == ["load_id", "predicted_rate"], "load_id,predicted_rate")
    add("Predictions: unique load_id", pred["load_id"].is_unique, "")
    add("Predictions: no missing load_id", not pred["load_id"].isna().any(), "")
    add("Predictions: no missing predicted_rate", not pred["predicted_rate"].isna().any(), "")
    add("Predictions: numeric", rates.notna().all(), "")
    add("Predictions: finite", bool(np.isfinite(rates).all()), "")
    add("Predictions: positive", bool((rates > 0).all()), f"min={rates.min():.4f}")
    add(
        "Predictions: IDs match template set",
        set(pred["load_id"].astype(str)) == set(tmpl["load_id"].astype(str)),
        "",
    )
    add(
        "Predictions: IDs match validation.csv set",
        set(pred["load_id"].astype(str)) == set(val["load_id"].astype(str)),
        "",
    )
    add(
        "Predictions: ordering matches template",
        pred["load_id"].astype(str).tolist() == tmpl["load_id"].astype(str).tolist(),
        "",
    )

    dec = pd.read_csv(root / "december-chart-inputs.csv")
    dec_ok = (
        len(dec) == 31
        and dec["pickup"].eq("Lexington").all()
        and dec["delivery"].eq("Fort Wayne").all()
        and np.isclose(dec["distance"], 360.0).all()
        and dec["equipment"].eq("Dry Van").all()
        and np.isclose(dec["weight"], 32000.0).all()
        and pd.to_datetime(dec["date"]).min().date().isoformat() == "2025-12-01"
        and pd.to_datetime(dec["date"]).max().date().isoformat() == "2025-12-31"
    )
    add("December inputs: fixed scenario", dec_ok, f"pred=${dec['predicted_rate'].iloc[0]:.2f}")

    chart = root / "scorer_results/candidate_december.png"
    add("December chart exists", chart.is_file(), str(chart))

    if report_docx.is_file():
        from docx import Document

        doc = Document(str(report_docx))
        text = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                text += "\n" + " | ".join(c.text for c in row.cells)
        img_count = sum(1 for r in doc.part.rels.values() if "image" in r.reltype)
        add("Report: embedded images", img_count >= 4, f"{img_count} images")
        add("Report: mentions 106.83", "106.83" in text, "")
        add("Report: mentions 113.29", "113.29" in text, "")
        add(
            "Report: dev MAE vs holdout distinction",
            ("development chronological validation" in text.lower() or "not nov" in text.lower())
            and ("no labels" in text.lower() or "no local ground truth" in text.lower()),
            "",
        )
        add("Report: score.py does not compute accuracy", "does not compute" in text.lower(), "")
        bad = ["TODO", "TBD", "insert chart", "final score:", "spotter score of"]
        add("Report: no placeholder/bad text", not any(b.lower() in text.lower() for b in bad), "")

    gi = root / ".gitignore"
    gi_text = gi.read_text(encoding="utf-8") if gi.is_file() else ""
    for pat in ["venv/", "__pycache__/", ".env"]:
        add(f".gitignore excludes {pat}", pat in gi_text, "")

    req_text = (root / "requirements.txt").read_text(encoding="utf-8")
    for pkg in ["matplotlib", "numpy", "pandas", "scikit-learn", "scipy"]:
        add(f"requirements includes {pkg}", pkg in req_text, "")
    for pkg in ["catboost", "xgboost", "lightgbm"]:
        add(f"requirements excludes {pkg}", pkg not in req_text.lower(), "")

    readme = (root / "README.md").read_text(encoding="utf-8")
    add("README: train_final.py command", "python scripts/train_final.py" in readme, "")
    add("README: score.py command", "python score.py --predictions validation_predictions.csv" in readme, "")
    add(
        "README: does not claim score.py MAE",
        "does not compute prediction accuracy" in readme.lower() or "does not compute accuracy" in readme.lower(),
        "",
    )

    failed = [r for r in results if r[1] == "FAIL"]
    print("FINAL_SUBMISSION_AUDIT")
    for req, status, evidence in results:
        ev = f" | {evidence}" if evidence else ""
        print(f"{status}: {req}{ev}")
    print(f"\nTOTAL: {len(results)} checks, {len(failed)} failed")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
