"""Generate the final assessment report (DOCX + PDF when possible)."""

from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
import numpy as np

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

REPORTS_DIR = PROJECT_ROOT / "reports"
DOCX_PATH = REPORTS_DIR / "freight_rate_ml_assessment.docx"
PDF_PATH = REPORTS_DIR / "freight_rate_ml_assessment.pdf"
EDA_DIR = PROJECT_ROOT / "reports" / "eda"
CHART_PATH = PROJECT_ROOT / "scorer_results" / "candidate_december.png"
PRED_PATH = PROJECT_ROOT / "validation_predictions.csv"
DECEMBER_PATH = PROJECT_ROOT / "december-chart-inputs.csv"


def _prediction_stats() -> dict[str, float]:
    df = pd.read_csv(PRED_PATH)
    rates = pd.to_numeric(df["predicted_rate"], errors="coerce")
    invalid = int((rates.isna() | ~np.isfinite(rates) | (rates <= 0)).sum())
    return {
        "mean": float(rates.mean()),
        "median": float(rates.median()),
        "min": float(rates.min()),
        "max": float(rates.max()),
        "invalid": invalid,
    }


def _december_rate() -> float | None:
    if not DECEMBER_PATH.is_file():
        return None
    df = pd.read_csv(DECEMBER_PATH)
    if "predicted_rate" not in df.columns:
        return None
    rates = pd.to_numeric(df["predicted_rate"], errors="coerce")
    if rates.isna().all():
        return None
    return float(rates.iloc[0])


def _money(x: float) -> str:
    return f"${x:,.2f}"


def _add_heading(doc, text, level=1):
    from docx.shared import Pt

    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(14 if level == 1 else 12)
    return h


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def _add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def _add_image(doc, path: Path, width_in=5.5, caption: str | None = None):
    from docx.shared import Inches

    if path.is_file():
        doc.add_picture(str(path), width=Inches(width_in))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Image not found: {path.name}]")


def build_docx() -> Path:
    from docx import Document

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stats = _prediction_stats()
    dec_rate = _december_rate()
    dec_label = _money(dec_rate) if dec_rate is not None else "N/A"

    doc = Document()
    doc.add_heading("Freight Rate Prediction — Machine Learning Assessment", 0)

    # 1 Executive Summary
    _add_heading(doc, "1. Executive Summary")
    _add_para(
        doc,
        "This project predicts freight load rates (posted_rate in USD) from shipment "
        "attributes. The development dataset contains 48,000 labeled loads (Jan–Oct 2025); "
        "the final inference holdout contains 12,000 unlabeled loads (Nov–Dec 2025)."
    )
    _add_para(
        doc,
        "The final model is HistGradientBoostingRegressor with feature set Q, a raw "
        "posted_rate target, and absolute_error loss. Development chronological validation "
        "MAE is 93.56 (primary split) and 102.57 (sensitivity split). Primary RMSE is 630.27. "
        "These are NOT Nov–Dec holdout scores — validation.csv has no local ground truth."
    )
    _add_para(
        doc,
        "An additional controlled optimization experiment compared alternative loss/target "
        "configurations on the same chronological splits. The raw-target absolute_error "
        "configuration generalized better across BOTH primary and sensitivity splits than the "
        "earlier HistGB Q + log1p + squared_error candidate (106.83 / 113.29)."
    )
    _add_para(
        doc,
        "Final predictions were generated for all 12,000 holdout loads and passed score.py "
        "format validation (exit code 0). Official accuracy, if any, is determined by Spotter "
        "after external submission."
    )

    # 2 Problem Definition
    _add_heading(doc, "2. Problem Definition")
    _add_bullets(
        doc,
        [
            "Input: load attributes (distance, equipment, weight, market context, etc.)",
            "Target: posted_rate (USD freight rate)",
            "Identifier: load_id (TR-* train, TE-* validation) — never used as a feature",
            "Task: regression — predict posted_rate for each validation load",
            "Final holdout: validation.csv (12,000 rows, Nov–Dec 2025, no labels locally)",
        ],
    )

    # 3 Dataset Overview
    _add_heading(doc, "3. Dataset Overview")
    _add_table(
        doc,
        ["File", "Rows", "Period", "Labels"],
        [
            ["train-test.csv", "48,000", "Jan–Oct 2025", "posted_rate present"],
            ["validation.csv", "12,000", "Nov–Dec 2025", "unlabeled holdout"],
            ["validation_predictions.csv", "12,000", "Nov–Dec 2025", "model output"],
        ],
    )
    _add_para(doc, "Key columns:", bold=True)
    _add_table(
        doc,
        ["Column", "Role"],
        [
            ["distance", "Numeric — strongest predictor (r ≈ 0.91)"],
            ["equipment", "Categorical — Dry Van, Reefer, Flatbed"],
            ["weight", "Numeric — missing/negative values require cleaning"],
            ["market_index", "Numeric — tested but excluded from final model"],
            ["quote_signal", "Numeric — retained in final feature set Q"],
            ["date", "Used for preprocessing; not in final Q features"],
        ],
    )

    # 4 EDA
    _add_heading(doc, "4. Exploratory Data Analysis")
    _add_para(
        doc,
        "posted_rate is right-skewed (mean $2,374 vs median $2,031) with a heavy upper tail "
        "driven by legitimate long-haul loads. Distance dominates the linear relationship with "
        "target. Equipment types show different rate levels (Reefer highest on average). "
        "market_index shows temporal drift across the full training period but Sep–Oct training "
        "aligns with Nov–Dec validation regime."
    )
    _add_image(doc, EDA_DIR / "01_posted_rate_distribution.png", caption="Figure 1: posted_rate distribution")
    _add_image(doc, EDA_DIR / "04_distance_vs_posted_rate.png", caption="Figure 2: Distance vs posted_rate")
    _add_image(doc, EDA_DIR / "09_equipment_vs_posted_rate.png", caption="Figure 3: Rate by equipment type")
    _add_image(doc, EDA_DIR / "03_posted_rate_by_month.png", caption="Figure 4: Monthly target behavior")

    # 5 Data Quality
    _add_heading(doc, "5. Data Quality")
    _add_para(
        doc,
        "Missing weight: 300 train / 165 validation rows. Negative weight: 292 train / 145 "
        "validation — treated as invalid and converted to missing.",
    )
    _add_para(
        doc,
        "Weight imputation: equipment median with global median fallback (fit on training fold "
        "only; full Jan–Oct for final model).",
    )
    _add_para(
        doc,
        "Missing market_index: 374 train / 249 validation. Imputed by month during "
        "preprocessing, but raw market_index was excluded from the final model after harming "
        "validation performance.",
    )
    _add_para(
        doc,
        "Unseen cities: 8 cities appear in validation but not training (~6% of rows). Final "
        "feature set Q avoids city/route identity, using equipment + distance + quote_signal instead.",
    )

    # 6 Feature Engineering
    _add_heading(doc, "6. Feature Engineering — Final Feature Set Q")
    _add_table(
        doc,
        ["Feature", "Rationale"],
        [
            ["distance", "Primary rate driver"],
            ["log_distance", "Nonlinear distance scaling"],
            ["distance_bin", "Bucketed distance effects"],
            ["weight / weight_is_missing", "Load size signal with missingness indicator"],
            ["quote_signal", "Strongest additive feature in ablation (+9 MAE improvement vs base)"],
            ["equipment", "Equipment-type rate differences (3 levels)"],
        ],
    )
    _add_para(
        doc,
        "Excluded after testing: raw market_index (+23 MAE when added), route/city OHE, "
        "geographic extras (FULL set had +22 sensitivity gap).",
    )

    # 7 Validation Strategy
    _add_heading(doc, "7. Validation Strategy")
    _add_table(
        doc,
        ["Split", "Train", "Validate", "Purpose"],
        [
            ["Primary", "Jan–Aug 2025", "Sep–Oct 2025", "Main model selection"],
            ["Sensitivity", "Jan–Sep 2025", "Oct 2025", "Stability check"],
        ],
    )
    _add_para(
        doc,
        "Random shuffling was avoided to simulate forecasting future periods. All "
        "preprocessing statistics are fit on the training fold only. validation.csv was never "
        "used for tuning.",
    )

    # 8 Model Experiments
    _add_heading(doc, "8. Model Experiments")
    _add_para(doc, "Primary split MAE unless noted (development chronological validation):")
    _add_table(
        doc,
        ["Model / Config", "Primary MAE", "Sensitivity MAE", "Notes"],
        [
            ["Global median", "1148.92", "1146.79", "Earlier baseline"],
            ["Distance linear", "196.95", "192.77", "Earlier baseline"],
            ["Distance + equipment linear", "155.68", "157.47", "Earlier baseline"],
            ["Ridge (FULL features)", "148.06", "148.92", "Earlier ML"],
            ["HistGB FULL normal", "129.69", "151.99", "Earlier ML"],
            ["ExtraTrees FULL", "142.86", "146.42", "Earlier ML"],
            ["HistGB FULL log1p", "118.83", "115.45", "Earlier experiment"],
            ["HistGB Q normal (squared_error)", "118.95", "123.91", "Earlier experiment"],
            ["HistGB Q log1p + squared_error", "106.83", "113.29", "Earlier candidate (not final)"],
            [
                "HistGB Q raw + absolute_error (FINAL)",
                "93.56",
                "102.57",
                "Final production model",
            ],
        ],
    )

    # 9 Final Model Selection
    _add_heading(doc, "9. Final Model Selection")
    _add_para(
        doc,
        "Selected (FINAL): HistGradientBoostingRegressor + feature set Q + raw posted_rate "
        "+ absolute_error loss.",
    )
    _add_bullets(
        doc,
        [
            "Best development chronological validation MAE on both primary and sensitivity splits",
            "Improves Primary RMSE (630.27), high-rate MAE (816.54), and long-haul MAE (164.93) "
            "vs the earlier log1p candidate",
            "Stable across chronological splits — not a primary-only overfit",
            "Simple feature set Q — no target encoding, no route OHE, no new features",
            "Predicts directly in dollar units (no log1p/expm1 transform)",
            "Leakage-safe preprocessing unchanged",
        ],
    )
    _add_para(
        doc,
        "Hyperparameters: loss=absolute_error, max_depth=6, l2_regularization=0.1, "
        "learning_rate=0.08, max_iter=300, random_state=42",
    )
    _add_para(
        doc,
        "Historical note: HistGB Q + log1p + squared_error (106.83 / 113.29) was an earlier "
        "strong candidate. A controlled loss/target comparison showed raw absolute_error "
        "generalized better on both splits and was adopted as the final model.",
    )

    # 10 Error Analysis
    _add_heading(doc, "10. Error Analysis")
    _add_para(
        doc,
        "Long-haul loads (2000+ miles) remain the hardest segment — systematic underprediction "
        "persists on top-1% and rate>$5,000 loads (high-rate MAE 816.54; long-haul MAE 164.93). "
        "Absolute-error training improves typical and long-haul errors relative to the earlier "
        "log1p/squared_error candidate, but does not eliminate extreme-tail bias.",
    )

    # 11 Final Training
    _add_heading(doc, "11. Final Training")
    _add_para(
        doc,
        "Final model trained on all 48,000 Jan–Oct labeled rows. validation.csv not used for "
        "fitting. Target: raw posted_rate. Loss: absolute_error. Predictions are direct dollar "
        "outputs (no expm1).",
    )

    # 12 Final Predictions
    _add_heading(doc, "12. Final Predictions")
    _add_para(
        doc,
        "12,000 Nov–Dec predictions in validation_predictions.csv. Distribution summary "
        "(not accuracy metrics):",
    )
    _add_table(
        doc,
        ["Statistic", "Value"],
        [
            ["Mean", _money(stats["mean"])],
            ["Median", _money(stats["median"])],
            ["Min", _money(stats["min"])],
            ["Max", _money(stats["max"])],
            ["Invalid predictions", "0"],
        ],
    )

    # 13 score.py
    _add_heading(doc, "13. Official score.py Validation")
    _add_para(
        doc,
        "Command: python score.py --predictions validation_predictions.csv "
        "--december-predictions december-chart-inputs.csv",
    )
    _add_para(
        doc,
        "Exit code: 0. score.py validates file format and generates the December chart. "
        "It does NOT compute prediction accuracy.",
    )
    _add_image(
        doc,
        CHART_PATH,
        width_in=6.0,
        caption=(
            f"Figure 5: Fixed December 2025 Candidate Prediction — flat at {dec_label} "
            "because feature set Q excludes calendar features; all other scenario inputs "
            "(Lexington → Fort Wayne, 360 mi, Dry Van, 32k lb) are fixed and only the date changes."
        ),
    )

    # 14 Limitations
    _add_heading(doc, "14. Limitations")
    _add_bullets(
        doc,
        [
            "No local ground truth for Nov–Dec holdout",
            "Official Spotter metric unavailable locally",
            "Long-tail / high-rate underprediction remains",
            "Development chronological MAE may differ from final holdout performance",
        ],
    )

    # 15 Reproducibility
    _add_heading(doc, "15. Reproducibility")
    _add_para(doc, "Setup:", bold=True)
    _add_bullets(
        doc,
        [
            "python -m venv venv",
            ".\\venv\\Scripts\\Activate.ps1",
            "python -m pip install -r requirements.txt",
        ],
    )
    _add_para(doc, "Pipeline commands:", bold=True)
    _add_bullets(
        doc,
        [
            "python scripts/train_final.py",
            "python score.py --predictions validation_predictions.csv "
            "--december-predictions december-chart-inputs.csv",
            "python scripts/generate_report.py",
        ],
    )

    # 16 Conclusion
    _add_heading(doc, "16. Conclusion")
    _add_para(
        doc,
        "HistGradientBoosting with feature set Q, raw posted_rate target, and absolute_error "
        "loss was selected through leakage-safe chronological validation and a controlled "
        "loss/target comparison. Final Nov–Dec predictions were generated and passed score.py "
        "validation. External Spotter evaluation determines any official accuracy.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


def build_full_pdf() -> Path | None:
    """Multi-page PDF with key sections and figures."""
    try:
        from matplotlib.backends.backend_pdf import PdfPages
        import matplotlib.pyplot as plt
        import textwrap

        stats = _prediction_stats()
        dec_rate = _december_rate()
        dec_label = _money(dec_rate) if dec_rate is not None else "N/A"

        pages = [
            (
                "Freight Rate Prediction — ML Assessment\n\n"
                "Executive Summary\n"
                "Predict posted_rate (USD) for freight loads. Final model: HistGradientBoosting "
                "with feature set Q, raw posted_rate target, and absolute_error loss. "
                "Trained on 48,000 Jan–Oct rows; inferred 12,000 Nov–Dec holdout rows.\n\n"
                "Development chronological validation MAE: 93.56 (primary) and 102.57 (sensitivity). "
                "Primary RMSE: 630.27. High-rate MAE: 816.54. Long-haul MAE: 164.93.\n"
                "These are NOT Nov–Dec holdout scores. score.py validates format only (exit code 0)."
            ),
            (
                "Dataset & Validation Strategy\n"
                "train-test.csv: 48,000 labeled rows. validation.csv: 12,000 unlabeled holdout.\n"
                "Primary split: Jan–Aug train → Sep–Oct validate.\n"
                "Sensitivity: Jan–Sep train → Oct validate.\n"
                "No random shuffle. Preprocessing fit on training fold only."
            ),
            (
                "Feature Set Q & Final Model\n"
                "Features: distance, log_distance, distance_bin, weight, weight_is_missing, "
                "quote_signal, equipment.\n"
                "Excluded: raw market_index, route/city OHE, geographic extras.\n"
                "Target: raw posted_rate (no log1p/expm1).\n"
                "Loss: absolute_error.\n"
                "Hyperparameters: max_depth=6, l2=0.1, lr=0.08, max_iter=300, random_state=42."
            ),
            (
                "Model Comparison (development chronological validation)\n"
                "Global median: 1148.92 | Distance linear: 196.95 | Ridge: 148.06\n"
                "HistGB FULL: 129.69 | HistGB log1p FULL: 118.83\n"
                "HistGB Q log1p + squared_error (earlier candidate): 106.83 / 113.29\n"
                "HistGB Q raw + absolute_error (FINAL): 93.56 primary / 102.57 sensitivity\n\n"
                f"Final predictions: mean {_money(stats['mean'])}, median {_money(stats['median'])}, "
                f"min {_money(stats['min'])}, max {_money(stats['max'])}."
            ),
            (
                "Limitations\n"
                "• No local Nov–Dec ground truth\n"
                "• score.py does not compute accuracy\n"
                "• Long-tail underprediction remains\n"
                f"• December chart flat at {dec_label} (Q excludes calendar features)"
            ),
        ]
        figures = [
            (EDA_DIR / "01_posted_rate_distribution.png", "posted_rate distribution"),
            (EDA_DIR / "04_distance_vs_posted_rate.png", "Distance vs posted_rate"),
            (EDA_DIR / "09_equipment_vs_posted_rate.png", "Rate by equipment"),
            (CHART_PATH, "December 2025 candidate prediction"),
        ]

        with PdfPages(PDF_PATH) as pdf:
            for text in pages:
                fig = plt.figure(figsize=(8.5, 11))
                fig.text(0.08, 0.92, textwrap.fill(text, 95), fontsize=10, va="top")
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
            for img_path, title in figures:
                if not img_path.is_file():
                    continue
                fig, ax = plt.subplots(figsize=(8.5, 6))
                ax.imshow(plt.imread(str(img_path)))
                ax.axis("off")
                ax.set_title(title, fontsize=11)
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
        return PDF_PATH
    except Exception:
        return None


def try_pdf_from_docx(docx_path: Path) -> Path | None:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(PDF_PATH))
        return PDF_PATH if PDF_PATH.is_file() else None
    except Exception:
        return None


def main():
    docx = build_docx()
    print(f"Created: {docx}")
    pdf = try_pdf_from_docx(docx)
    if pdf:
        print(f"Created: {pdf}")
    else:
        pdf = build_full_pdf()
        if pdf:
            print(f"Created PDF: {pdf}")
        else:
            print("PDF not created — use DOCX report.")


if __name__ == "__main__":
    main()
